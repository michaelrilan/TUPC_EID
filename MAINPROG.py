from login import *
from register import *
from access import *
from update import *
import re
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import QIntValidator,QRegExpValidator
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
import mysql.connector
from mysql.connector import errorcode
import tempfile
import sys
from os import listdir
from os.path import isfile, join
from os.path import dirname
import os
from pathlib import Path
from os.path import expanduser
from PyQt5.QtWidgets import (QApplication,QWidget)
from PyQt5.QtCore import QRect,QPropertyAnimation
import datetime
from datetime import date
from datetime import datetime,timedelta
from dateutil.relativedelta import relativedelta
import cv2
import numpy as np
from PyQt5.QtGui import QImage
from PyQt5.QtGui import QPixmap
from PyQt5.QtCore import QTimer
import uuid
import time
from PIL import Image
import PIL
import glob
import docx
from docx import Document
from docx.shared import Inches
cap = cv2.VideoCapture(0)
conn = mysql.connector.connect(host="localhost", user="root", passwd="admin", database="employeeinfodb")
firstID = None
wawawa = []
class newWindow1(Ui_MainWindow1,QtWidgets.QMainWindow):
    def __init__(self):
        super(newWindow1, self).__init__()
        self.setupUi(self)
        self.pushButton.clicked.connect(self.nextaccess)
        self.pushButton2.clicked.connect(self.nextregister)
        self.threedays = date.today() + timedelta(3)
        self.twodays = date.today() + timedelta(2)
        self.oneday = date.today() + timedelta(1)
        self.sub_3years3days = self.threedays + relativedelta(years=-3)
        self.sub_3years2days = self.twodays + relativedelta(years=-3)
        self.sub_3years1day = self.oneday + relativedelta(years=-3)
        print(self.sub_3years3days)
        print(self.sub_3years2days)
        print(self.sub_3years1day)
        self.cursornotif = conn.cursor()
        self.querynotif = '''SELECT LASTNAME,FIRSTNAME,MIDDLENAME,DATE_OF_LAST_PROMOTION 
        FROM EMPLOYEE_INFO WHERE DATE_OF_LAST_PROMOTION = %s or DATE_OF_LAST_PROMOTION = %s or DATE_OF_LAST_PROMOTION = %s;
        '''
        self.datanotif=(self.sub_3years3days,self.sub_3years2days,self.sub_3years1day)

        self.cursornotif.execute(self.querynotif,self.datanotif)
        self.fetchnotif = self.cursornotif.fetchall()
        print(self.fetchnotif)
        self.setFixedSize(381, 418)
    def returnfetchednotif(self):
        return self.fetchnotif
    def nextaccess(self):
        user=self.lineusername.text()
        pw=self.linepassword.text()
        currlog=conn.cursor()
        currlog1=conn.cursor()
        datee = date.today()
        datesss = str(datee)
        print('waaaaaaaaaaaaaaaa '+str(datee))
        x = datesss.split("-")
        print(x)
        querylog="SELECT username,passwrd from account;"
        currlog.execute(querylog)
        a=currlog.fetchall()
        leng=len(a)
        print(a)
        if a is not None:
            for iii in range(len(a)):
                if a[iii][0] == user and a[iii][1] == pw:
                    w3.show()

                    w1.hide()
                    w2.hide()
                    currdate = conn.cursor()
                    currdate.execute("select date_registered from account where username='" + user+"';")
                    a1 = currdate.fetchone()
                    w3.lineEdit_8.setText(str(user))
                    w3.lineEdit_9.setText(str(a1[0]))
                    w3.frameinfo.show()
                    w3.frameuser.hide()
                    w3.framepass.hide()
                    w3.framepass_2.hide()

                    querylog1="SELECT LASTNAME,FIRSTNAME,MIDDLENAME,DAY_OF_PROMOTION from employee_info;"
                    currlog1.execute(querylog1)
                    a1 =currlog1.fetchall()
                    print(a1)

                    pro_list = []
                    for ii in range(len(a1)):
                        date1 = a1[ii][3]
                        date2 = str(date1)
                        splt = date2.split('/')
                        if x[0] == splt[0]:
                            pro_list.append([a1[ii][0],a1[ii][1],a1[ii][2],a1[ii][3]])
                    print('waaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
                    print(pro_list)




                    if len(pro_list):
                        w3.listnotif.clear()
                        
                        for i in range(len(pro_list)):
                            '''nextprom = self.fetchnotif[i][3]+ relativedelta(years=3)
                            resultforprompt=abs(nextprom-date.today()).days
                            print(resultforprompt)'''
                            w3.listnotif.addItem(str(pro_list[i][0]) + "," + str(pro_list[i][1]) + "," + str(
                                pro_list[i][2]))
                    else:
                        w3.listnotif.clear()
                        w3.listnotif.addItem("No notification so far")
                    return
            else:
                QMessageBox.warning(self,"Error","Invalid Credentials")

        else:
            QMessageBox.warning(self,"Error","Invalid Credentials")

    def nextregister(self):
        w3.hide()
        w1.hide()
        w2.show()

class newWindow2(Ui_MainWindow2,QtWidgets.QMainWindow):
    def __init__(self):
        super(newWindow2, self).__init__()
        self.setupUi(self)
        self.pushButton.clicked.connect(self.backtologin)
        self.pushButton_2.clicked.connect(self.registerr)
        self.setFixedSize(390, 483)


    def backtologin(self):
        w1.show()
        w2.hide()
        w3.hide()
        self.lineEdit.clear()
        self.linepassword1.clear()
        self.linepassword2.clear()
        self.lineverify.clear()
    def registerr(self):
        user = self.lineEdit.text()
        pw1=self.linepassword1.text()
        pw2=self.linepassword2.text()
        verif=self.lineverify.text()
        if pw1 == pw2:
            cur=conn.cursor()

            cur.execute("select * from codee;")
            a=cur.fetchone()
            if verif == str(a[0]):
                curreg=conn.cursor()
                queryreg="insert into account (username,passwrd,date_registered)values(%s,%s,curdate());"
                data=(user,pw2)
                curreg.execute(queryreg,data)
                conn.commit()
                QMessageBox.about(self,"Registered","You can now access the system ")
                self.lineEdit.clear()
                self.linepassword1.clear()
                self.linepassword2.clear()
                self.lineverify.clear()
            else:
                QMessageBox.warning("Error","Invalid code")
        else:
            QMessageBox.warning(self,"Error","Passwords unequal")


class newWindow3(Ui_MainWindow3,QtWidgets.QMainWindow):
    def __init__(self):
        super(newWindow3, self).__init__()
        self.setupUi(self)

        self.pic.hide()
        self.textData.hide()
        self.pushPrint.hide()
        self.pushModify.hide()
        self.lineEdit_8.setReadOnly(True)
        self.lineEdit_9.setReadOnly(True)
        self.label_77.setDisabled(True)
        self.textEdit.setDisabled(True)
        self.textviewnotif.hide()
        self.pushupdatenotif.hide()
        self.frame_3.hide()
        self.frameinfo.show()
        self.frameuser.hide()
        self.framepass.hide()
        self.framepass_2.hide()
        self.frameList.setGeometry(QtCore.QRect(0, 88, 931, 481))
        self.pushModify.clicked.connect(self.showmodify)
        self.setFixedWidth(931)
        self.setFixedHeight(630)
        self.today = date.today()
        self.d = QDate(self.today)
        self.dateEdit.setDate(self.d)
        self.dateEdit_2.setDate(self.d)
        self.dateEdit_3.setDate(self.d)
        self.listE.setSortingEnabled(True)

        self.path = dirname(__file__)
# side bar//////////////////////////////////////////////////////////////////////////
        self.pushmenuin.clicked.connect(self.menuin)
        self.pushmenuout.clicked.connect(self.menuhide)
        self.commandLinkButton_4.clicked.connect(self.showhome)
        self.commandLinkButton.clicked.connect(self.showadd)
        self.pushACC.clicked.connect(self.showaccount)
        self.pushACC_2.clicked.connect(self.shownotif)
#   account settings //////////////////////////////////////////////////////////////////////////
        self.pushButton.clicked.connect(self.showAccUser)
        self.pushButton_10.clicked.connect(self.showAccInfo)
        self.pushButton_8.clicked.connect(self.showAccPass)
        self.pushButton_9.clicked.connect(self.showAccVerify)
#//////////////////////////////////////////////////////////////////////////
        self. confirmuser.clicked.connect(self.changeuser)
        self.confirmpass.clicked.connect(self.changepassword)
        self.confirmpass_3.clicked.connect(self.changevcode)
        self.comboBox_3.currentTextChanged.connect(self.showremarks)
#/////////////////////////////////////////////////////////////////////////
        self.comboList.currentTextChanged.connect(self.showlistemploy)
        self.pushAdd.clicked.connect(self.addemploy)
        self.listE.itemClicked.connect(self.clicklistE)
        self.listnotif.itemClicked.connect(self.clicklistnotif)
        self.pushupdatenotif.clicked.connect(self.updatenotif)
        self.lineSearch.textChanged.connect(self.searchlist)
        self.pushPrint.clicked.connect(self.printinfoselectedemployee)
        self.pushLogout.clicked.connect(self.logout)
        self.pushbrowsepic.clicked.connect(self.addemploybrowse)
        self.pushPrint_2.clicked.connect(self.printlistpersonels)
        #~~~~~~~~~~~~~~~~~~~~~~~~~~~VALIDATOR~~~~~~~~~~~~~~~~~~~~~~~~~~
        regex = QtCore.QRegExp("^[a-zA-Z\s]*$")
        regnumnine = QtCore.QRegExp("^(\d{9})$")
        regnumten = QtCore.QRegExp("^(\d{10})$")
        regnumeleven = QtCore.QRegExp("^(\d{11})$")
        regnumtwelve = QtCore.QRegExp("^(\d{12})$")
        validatornine = QtGui.QRegExpValidator(regnumnine)
        validatorten = QtGui.QRegExpValidator(regnumten)
        validatoreleven = QtGui.QRegExpValidator(regnumeleven)
        validatortwelve = QtGui.QRegExpValidator(regnumtwelve)
        validator = QtGui.QRegExpValidator(regex)

        self.lineFirstname_3.setValidator(validator)
        self.lineMid_3.setValidator(validator)
        self.lineLastname_3.setValidator(validator)
        self.lineDesignation_3.setValidator(validator)
        self.lineStepincre_3.setValidator(QIntValidator())
        self.linePhone_9.setValidator(QIntValidator())
        self.lineSalaryg_3.setValidator(QIntValidator())
        self.lineStepincre_3.setValidator(QIntValidator())
        self.linePhone_8.setValidator(validatornine)
        self.browsedpic = None
        self.linePhone_7.setValidator(validatoreleven)
        self.lineTin_3.setValidator(validatornine)
        self.lineGsis_3.setValidator(validatoreleven)
        self.lineSss_3.setValidator(validatorten)
        self.linePagibig_3.setValidator(validatortwelve)
        self.linePhil_3.setValidator(validatortwelve)
        #Error HANDLING~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        self.lineFirstname_3.textChanged.connect(self.auto_capitalfirstname)
        self.lineMid_3.textChanged.connect(self.auto_capitalmiddlename)
        self.lineLastname_3.textChanged.connect(self.auto_capitallastname)
        self.lineDesignation_3.textChanged.connect(self.auto_capitaldesignation)

        self.labe = None
        self.ID = None
        #self.firstt = None
        #self.midd= None
        self.genders =None
        #self.existingba = None
        self.listE.clear()
        self.textData.setReadOnly(True)
        cur = conn.cursor()
        cur.execute(
            "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
        a = cur.fetchall()
        for i in range(len(a)):
            self.listE.addItem(str(a[i][0])+","+ "  " + str(a[i][1]) + "  " + str(a[i][2]))


#~~~~~~~~~~~~~~~~~~~~~~~~~~~~TRANSITIONS~~~~~~~~~~~~~~~~~

    def menuhide(self):
        self.pushmenuin.setDisabled(False)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(0, 85, 191, 545,))
        self.anim.setEndValue(QRect(-190,90,191,545))
        self.anim.start()

        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(190, 570, 931, 61))
        self.animfooter.setEndValue(QRect(0, 570, 931, 61))
        self.animfooter.start()

        if self.frameList.geometry() == QRect(190, 88, 931, 481):
            self.frameList.setDisabled(False)
            self.animframeList =  QPropertyAnimation(self.frameList, b"geometry")
            self.animframeList.setDuration(230)
            self.animframeList.setStartValue(QRect(190, 88, 931, 481))
            self.animframeList.setEndValue(QRect(0,88,931,481))
            self.animframeList.start()
        elif self.frameAdd_2.geometry() == QRect(190, 88, 931, 481):
            self.frameAdd_2.setDisabled(False)
            self.animframeAdd_2 = QPropertyAnimation(self.frameAdd_2, b"geometry")
            self.animframeAdd_2.setDuration(230)
            self.animframeAdd_2.setStartValue(QRect(190, 88, 931, 481))
            self.animframeAdd_2.setEndValue(QRect(0, 88, 931, 481))
            self.animframeAdd_2.start()

        elif self.frameACC.geometry() == QRect(190, 88, 931, 481):
            self.frameACC.setDisabled(False)
            self.animframeACC = QPropertyAnimation(self.frameACC, b"geometry")
            self.animframeACC.setDuration(230)
            self.animframeACC.setStartValue(QRect(190, 88, 931, 481))
            self.animframeACC.setEndValue(QRect(0, 88, 931, 481))
            self.animframeACC.start()

        elif self.frameleave.geometry() == QRect(190, 88, 931, 481):
            self.frameleave.setDisabled(False)
            self.animframeleave = QPropertyAnimation(self.frameleave, b"geometry")
            self.animframeleave.setDuration(230)
            self.animframeleave.setStartValue(QRect(190, 88, 931, 481))
            self.animframeleave.setEndValue(QRect(0, 88, 931, 481))
            self.animframeleave.start()

        elif self.framenotif.geometry() == QRect(190, 88, 931, 481):
            self.framenotif.setDisabled(False)
            self.animframenotif = QPropertyAnimation(self.framenotif, b"geometry")
            self.animframenotif.setDuration(230)
            self.animframenotif.setStartValue(QRect(190, 88, 931, 481))
            self.animframenotif.setEndValue(QRect(0, 88, 931, 481))
            self.animframenotif.start()

    def menuin(self):
        self.pushmenuin.setDisabled(True)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(-190,90,191,545))
        self.anim.setEndValue(QRect(0, 85, 191, 545))
        self.anim.start()
        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(0, 570, 931, 61))
        self.animfooter.setEndValue(QRect(190, 570, 931, 61))
        self.animfooter.start()
        if self.frameList.geometry() == QRect(0, 88, 931, 481):
            self.frameList.setDisabled(True)
            self.animframeList =  QPropertyAnimation(self.frameList, b"geometry")
            self.animframeList.setDuration(230)
            self.animframeList.setStartValue(QRect(0, 88, 931, 481))
            self.animframeList.setEndValue(QRect(190,88,931,481))
            self.animframeList.start()
        elif self.frameAdd_2.geometry() == QRect(0, 88, 931, 481):
            self.frameAdd_2.setDisabled(True)
            self.animframeAdd_2 = QPropertyAnimation(self.frameAdd_2, b"geometry")
            self.animframeAdd_2.setDuration(230)
            self.animframeAdd_2.setStartValue(QRect(0, 88, 931, 481))
            self.animframeAdd_2.setEndValue(QRect(190, 88, 931, 481))
            self.animframeAdd_2.start()

        elif self.frameACC.geometry() == QRect(0, 88, 931, 481):
            self.frameACC.setDisabled(True)
            self.animframeACC = QPropertyAnimation(self.frameACC, b"geometry")
            self.animframeACC.setDuration(230)
            self.animframeACC.setStartValue(QRect(0, 88, 931, 481))
            self.animframeACC.setEndValue(QRect(190, 88, 931, 481))
            self.animframeACC.start()

        elif self.frameleave.geometry() == QRect(0, 88, 931, 481):
            self.frameleave.setDisabled(True)
            self.animframeleave = QPropertyAnimation(self.frameleave, b"geometry")
            self.animframeleave.setDuration(230)
            self.animframeleave.setStartValue(QRect(0, 88, 931, 481))
            self.animframeleave.setEndValue(QRect(190, 88, 931, 481))
            self.animframeleave.start()

        elif self.framenotif.geometry() == QRect(0, 88, 931, 481):
            self.framenotif.setDisabled(True)
            self.animframenotif = QPropertyAnimation(self.framenotif, b"geometry")
            self.animframenotif.setDuration(230)
            self.animframenotif.setStartValue(QRect(0, 88, 931, 481))
            self.animframenotif.setEndValue(QRect(190, 88, 931, 481))
            self.animframenotif.start()
    
    def printlistpersonels(self):
        doc = docx.Document()
        employtype = self.comboList.currentText()
        if employtype == 'TEACHING PERSONNEL':
            cur = conn.cursor()
            cur.execute("select LASTNAME,FIRSTNAME,MIDDLENAME from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            doc.add_heading('LIST OF TEACHING PERSONNEL PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')
        elif employtype == 'NON-TEACHING PERSONNEL':
            cur = conn.cursor()
            cur.execute("select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            doc.add_heading('LIST OF NON-TEACHING PERSONNEL PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')
        elif employtype == 'PART TIME':
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            doc.add_heading('LIST OF PART TIME PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')
        elif employtype == 'CONTRACT OF SERVICE':
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            doc.add_heading('LIST OF CONTRACT OF SERVICE PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')

        elif employtype == 'SEPARATED':
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            doc.add_heading('LIST OF SEPARATED PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')
        else:
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
            a = cur.fetchall()
            doc.add_heading('LIST OF ALL PERSONELS',0).bold=True
            doc.add_heading('FULL NAMES',1).bold=True
            for i in range(len(a)):
                doc.add_paragraph(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
            doc.save('list.docx')
            os.system('start list.docx')
        
        #doc.add_picture('9874.png',width=Inches(1.1))
        



    def showhome(self):
        self.pushmenuin.setDisabled(False)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(0, 85, 191, 545, ))
        self.anim.setEndValue(QRect(-190, 90, 191, 545))
        self.anim.start()

        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(190, 570, 931, 61))
        self.animfooter.setEndValue(QRect(0, 570, 931, 61))
        self.animfooter.start()
        if self.frameList.geometry() == QRect(190, 88, 931, 481):
            self.frameList.setDisabled(False)
            self.frameAdd_2.setDisabled(False)
            self.frameACC.setDisabled(False)
            self.frameleave.setDisabled(False)
            self.framenotif.setDisabled(False)
            self.animframeList =  QPropertyAnimation(self.frameList, b"geometry")
            self.animframeList.setDuration(230)
            self.animframeList.setStartValue(QRect(190, 88, 931, 481))
            self.animframeList.setEndValue(QRect(0,88,931,481))
            self.animframeList.start()
        else:
            if self.frameAdd_2.geometry() == QRect(190, 88, 931, 481):
                self.AddOut()
                self.homein()

            elif self.frameACC.geometry() == QRect(190, 88, 931, 481):
                self.ACCout()
                self.homein()
            elif self.frameleave.geometry() == QRect(190, 88, 931, 481):
                self.leaveout()
                self.homein()
            elif self.framenotif.geometry() == QRect(190, 88, 931, 481):
                self.notifout()
                self.homein()


    def showadd(self):
        self.pushmenuin.setDisabled(False)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(0, 85, 191, 545, ))
        self.anim.setEndValue(QRect(-190, 90, 191, 545))
        self.anim.start()

        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(190, 570, 931, 61))
        self.animfooter.setEndValue(QRect(0, 570, 931, 61))
        self.animfooter.start()

        if self.frameAdd_2.geometry() == QRect(190, 88, 931, 481):
            self.frameList.setDisabled(False)
            self.frameAdd_2.setDisabled(False)
            self.frameACC.setDisabled(False)
            self.frameleave.setDisabled(False)
            self.framenotif.setDisabled(False)
            self.animframeAdd_2 =  QPropertyAnimation(self.frameAdd_2, b"geometry")
            self.animframeAdd_2.setDuration(230)
            self.animframeAdd_2.setStartValue(QRect(190, 88, 931, 481))
            self.animframeAdd_2.setEndValue(QRect(0,88,931,481))
            self.animframeAdd_2.start()
        else:
            if self.frameList.geometry() == QRect(190, 88, 931, 481):
                self.Addin()
                self.homeout()
            elif self.frameACC.geometry() == QRect(190, 88, 931, 481):
                self.ACCout()
                self.Addin()
            elif self.frameleave.geometry() == QRect(190, 88, 931, 481):
                self.leaveout()
                self.Addin()
            elif self.framenotif.geometry() == QRect(190, 88, 931, 481):
                self.notifout()
                self.Addin()
    def showaccount(self):
        self.pushmenuin.setDisabled(False)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(0, 85, 191, 545, ))
        self.anim.setEndValue(QRect(-190, 90, 191, 545))
        self.anim.start()

        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(190, 570, 931, 61))
        self.animfooter.setEndValue(QRect(0, 570, 931, 61))
        self.animfooter.start()

        if self.frameACC.geometry() == QRect(190, 88, 931, 481):
            self.frameList.setDisabled(False)
            self.frameAdd_2.setDisabled(False)
            self.frameACC.setDisabled(False)
            self.frameleave.setDisabled(False)
            self.framenotif.setDisabled(False)
            self.animframeACC = QPropertyAnimation(self.frameACC, b"geometry")
            self.animframeACC.setDuration(230)
            self.animframeACC.setStartValue(QRect(190, 88, 931, 481))
            self.animframeACC.setEndValue(QRect(0, 88, 931, 481))
            self.animframeACC.start()
        else:
            if self.frameList.geometry() == QRect(190, 88, 931, 481):
                self.ACCin()
                self.homeout()
            elif self.frameleave.geometry() == QRect(190, 88, 931, 481):
                self.ACCin()
                self.leaveout()
            elif self.frameAdd_2.geometry() == QRect(190, 88, 931, 481):
                self.ACCin()
                self.AddOut()
            elif self.framenotif.geometry() == QRect(190, 88, 931, 481):
                self.notifout()
                self.ACCin()
    def shownotif(self):
        self.pushmenuin.setDisabled(False)
        self.anim = QPropertyAnimation(self.menu, b"geometry")
        self.anim.setDuration(230)
        self.anim.setStartValue(QRect(0, 85, 191, 545, ))
        self.anim.setEndValue(QRect(-190, 90, 191, 545))
        self.anim.start()

        self.animfooter = QPropertyAnimation(self.footer, b"geometry")
        self.animfooter.setDuration(230)
        self.animfooter.setStartValue(QRect(190, 570, 931, 61))
        self.animfooter.setEndValue(QRect(0, 570, 931, 61))
        self.animfooter.start()

        if self.framenotif.geometry() == QRect(190, 88, 931, 481):
            self.frameList.setDisabled(False)
            self.frameAdd_2.setDisabled(False)
            self.frameACC.setDisabled(False)
            self.frameleave.setDisabled(False)
            self.framenotif.setDisabled(False)
            self.animframenotif = QPropertyAnimation(self.framenotif, b"geometry")
            self.animframenotif.setDuration(230)
            self.animframenotif.setStartValue(QRect(190, 88, 931, 481))
            self.animframenotif.setEndValue(QRect(0, 88, 931, 481))
            self.animframenotif.start()
        else:
            if self.frameList.geometry() == QRect(190, 88, 931, 481):
                self.notifin()
                self.homeout()
            elif self.frameleave.geometry() == QRect(190, 88, 931, 481):
                self.notifin()
                self.leaveout()
            elif self.frameAdd_2.geometry() == QRect(190, 88, 931, 481):
                self.notifin()
                self.AddOut()
            elif self.frameACC.geometry() == QRect(190, 88, 931, 481):
                self.notifin()
                self.ACCout()
    def homeout(self):
        self.frameList.setDisabled(False)
        self.frameAdd_2.setDisabled(False)
        self.frameACC.setDisabled(False)
        self.frameleave.setDisabled(False)
        self.framenotif.setDisabled(False)

        self.animframeList = QPropertyAnimation(self.frameList, b"geometry")
        self.animframeList.setDuration(230)
        self.animframeList.setStartValue(QRect(190, 88, 931, 481))
        self.animframeList.setEndValue(QRect(1010, 88, 931, 481))
        self.animframeList.start()
    def AddOut(self):
        self.frameList.setDisabled(False)
        self.frameAdd_2.setDisabled(False)
        self.frameACC.setDisabled(False)
        self.frameleave.setDisabled(False)
        self.framenotif.setDisabled(False)
        self.animframeAdd_2 = QPropertyAnimation(self.frameAdd_2, b"geometry")
        self.animframeAdd_2.setDuration(230)
        self.animframeAdd_2.setStartValue(QRect(190, 88, 931, 481))
        self.animframeAdd_2.setEndValue(QRect(1010, 88, 931, 481))
        self.animframeAdd_2.start()
    def ACCout(self):
        self.frameList.setDisabled(False)
        self.frameAdd_2.setDisabled(False)
        self.frameACC.setDisabled(False)
        self.frameleave.setDisabled(False)
        self.framenotif.setDisabled(False)
        self.animframeACC = QPropertyAnimation(self.frameACC, b"geometry")
        self.animframeACC.setDuration(230)
        self.animframeACC.setStartValue(QRect(190, 88, 931, 481))
        self.animframeACC.setEndValue(QRect(1010, 88, 931, 481))
        self.animframeACC.start()
    def leaveout(self):
        self.frameList.setDisabled(False)
        self.frameAdd_2.setDisabled(False)
        self.frameACC.setDisabled(False)
        self.frameleave.setDisabled(False)
        self.framenotif.setDisabled(False)
        self.animframeleave = QPropertyAnimation(self.frameleave, b"geometry")
        self.animframeleave.setDuration(230)
        self.animframeleave.setStartValue(QRect(190, 88, 931, 481))
        self.animframeleave.setEndValue(QRect(1010, 88, 931, 481))
        self.animframeleave.start()
    def notifout(self):
        self.frameList.setDisabled(False)
        self.frameAdd_2.setDisabled(False)
        self.frameACC.setDisabled(False)
        self.frameleave.setDisabled(False)
        self.framenotif.setDisabled(False)
        self.animframenotif = QPropertyAnimation(self.framenotif, b"geometry")
        self.animframenotif.setDuration(230)
        self.animframenotif.setStartValue(QRect(190, 88, 931, 481))
        self.animframenotif.setEndValue(QRect(1010, 88, 931, 481))
        self.animframenotif.start()

    def homein(self):
        self.frameList.setDisabled(False)
        self.animframeList = QPropertyAnimation(self.frameList, b"geometry")
        self.animframeList.setDuration(230)
        self.animframeList.setStartValue(QRect(1010, 88, 931, 481))
        self.animframeList.setEndValue(QRect(0, 88, 931, 481))
        self.animframeList.start()
    def Addin(self):
        self.frameAdd_2.setDisabled(False)
        self.animframeAdd_2 = QPropertyAnimation(self.frameAdd_2, b"geometry")
        self.animframeAdd_2.setDuration(230)
        self.animframeAdd_2.setStartValue(QRect(1010, 88, 931, 481))
        self.animframeAdd_2.setEndValue(QRect(0, 88, 931, 481))
        self.animframeAdd_2.start()
    def ACCin(self):
        self.frameACC.setDisabled(False)
        self.animframeACC = QPropertyAnimation(self.frameACC, b"geometry")
        self.animframeACC.setDuration(230)
        self.animframeACC.setStartValue(QRect(1010, 88, 931, 481))
        self.animframeACC.setEndValue(QRect(0, 88, 931, 481))
        self.animframeACC.start()

    def leavein(self):
        self.frameleave.setDisabled(False)
        self.animframeleave = QPropertyAnimation(self.frameleave, b"geometry")
        self.animframeleave.setDuration(230)
        self.animframeleave.setStartValue(QRect(1010, 88, 931, 481))
        self.animframeleave.setEndValue(QRect(0, 88, 931, 481))
        self.animframeleave.start()
    def notifin(self):
        self.framenotif.setDisabled(False)
        self.animframenotif = QPropertyAnimation(self.framenotif, b"geometry")
        self.animframenotif.setDuration(230)
        self.animframenotif.setStartValue(QRect(1010, 88, 931, 481))
        self.animframenotif.setEndValue(QRect(0, 88, 931, 481))
        self.animframenotif.start()
    def showAccInfo(self):

        self.frameinfo.show()
        self.framepass.hide()
        self.framepass_2.hide()
        self.frameuser.hide()
    def showAccUser(self):
        self.frameinfo.hide()
        self.framepass.hide()
        self.framepass_2.hide()
        self.frameuser.show()
    def showAccPass(self):
        self.frameinfo.hide()
        self.framepass.show()
        self.framepass_2.hide()
        self.frameuser.hide()
    def showAccVerify(self):
        self.frameinfo.hide()
        self.framepass.hide()
        self.framepass_2.show()
        self.frameuser.hide()
    def logout(self):
        ret = QMessageBox.question(self, 'Signing Out?', "Are you sure you want to logout?",
                                   QMessageBox.Yes | QMessageBox.No)

        if ret == QMessageBox.Yes:
            w1.show()
            w3.hide()
            w2.hide()
        self.lineEdit_8.clear()
        self.lineEdit_9.clear()

        self.lineEdit_11.clear()
        self.lineEdit_12.clear()
        self.lineEdit_13.clear()

        self.lineEdit_14.clear()
        self.lineEdit_15.clear()
        self.lineEdit_16.clear()

        self.lineEdit_20.clear()
        self.lineEdit_21.clear()
        self.lineEdit_22.clear()
        w1.lineusername.clear()
        w1.linepassword.clear()
    def changeuser(self):
        newuser=self.lineEdit_11.text()
        currentuser = w1.lineusername.text()
        confirmvalid=self.lineEdit_12.text()
        confirmpass=self.lineEdit_13.text()
        pass1= w1.linepassword.text()

        #confirm validation code
        curvalid=conn.cursor()
        queryvalid="SELECT * FROM codee;"
        curvalid.execute(queryvalid)
        getvalid=curvalid.fetchone()


        if getvalid[0]==confirmvalid and pass1 == confirmpass:
            curuser = conn.cursor()
            queryuser = "UPDATE ACCOUNT SET USERNAME = %s where username = %s;"
            datauser = (newuser,currentuser)
            curuser.execute(queryuser,datauser)
            conn.commit()
            QMessageBox.about(self,"Changed","Your Username has been Changed, please Re-Login to Restart")
            self.lineEdit_11.clear()
            self.lineEdit_12.clear()
            self.lineEdit_13.clear()

            self.lineEdit_14.clear()
            self.lineEdit_15.clear()
            self.lineEdit_16.clear()

            self.lineEdit_20.clear()
            self.lineEdit_21.clear()
            self.lineEdit_22.clear()
            w1.lineusername.clear()
            w1.linepassword.clear()
            w1.show()
            w2.hide()
            w3.hide()
        else:
            QMessageBox.warning(self,"Invalid Credentials","Please check your inputted data")
    def changepassword(self):
        currpasss = self.lineEdit_14.text()
        newpasss = self.lineEdit_15.text()
        confirmpasss = self.lineEdit_16.text()
        userlog = w1.lineusername.text()
        passlog = w1.linepassword.text()

        if currpasss == passlog:
            if confirmpasss == newpasss:
                currchangepass = conn.cursor()
                querycpass = "update account set passwrd = %s where username = %s and passwrd = %s;"
                data = (confirmpasss,userlog,currpasss)
                currchangepass.execute(querycpass,data)
                conn.commit()
                QMessageBox.about(self,"Changed","Your Password has been Changed, please Re-Login to Restart")
                self.lineEdit_11.clear()
                self.lineEdit_12.clear()
                self.lineEdit_13.clear()

                self.lineEdit_14.clear()
                self.lineEdit_15.clear()
                self.lineEdit_16.clear()

                self.lineEdit_20.clear()
                self.lineEdit_21.clear()
                self.lineEdit_22.clear()
                w1.lineusername.clear()
                w1.linepassword.clear()
                w1.show()
                w2.hide()
                w3.hide()
            else:
                QMessageBox.warning(self, "Error", "New and Confirm Password Incorrect")
        else:
            QMessageBox.warning(self,"Error","Incorrect Password")
    def changevcode(self):
        currvalid = self.lineEdit_20.text()
        newvalid = self.lineEdit_21.text()
        confirmvalid = self.lineEdit_22.text()

        cursorvalid = conn.cursor()
        queryvalid = "SELECT * FROM CODEE;"
        cursorvalid.execute(queryvalid)
        a = cursorvalid.fetchone()

        if currvalid == a[0]:
            if newvalid == confirmvalid:
                queryvalid1 = "update codee set verificationcode = %s where verificationcode = %s;"
                data1 = (str(confirmvalid),str(currvalid))
                cursor = conn.cursor()
                cursor.execute(queryvalid1 , data1)
                conn.commit()
                QMessageBox.about(self,"Error","Your Verification Code has been changed, Please Re-Login to restart")
                self.lineEdit_11.clear()
                self.lineEdit_12.clear()
                self.lineEdit_13.clear()

                self.lineEdit_14.clear()
                self.lineEdit_15.clear()
                self.lineEdit_16.clear()

                self.lineEdit_20.clear()
                self.lineEdit_21.clear()
                self.lineEdit_22.clear()
                w1.lineusername.clear()
                w1.linepassword.clear()
                w1.show()
                w2.hide()
                w3.hide()
            else:
                QMessageBox.warning(self, "Error", "New and Confirm Verification Code is not equal")
        else:
            QMessageBox.warning(self,"Error","Incorrect Verification Code")
    def showremarks(self):
        emp = self.comboBox_3.currentText()
        if emp == "SEPARATED":
            self.label_77.setDisabled(False)
            self.textEdit.setDisabled(False)
        else:
            self.label_77.setDisabled(True)
            self.textEdit.setDisabled(True)
            self.textEdit.clear()
    def showmodify(self):
        w4.show()
        self.hide()
        cursor = conn.cursor()
        query = "SELECT * FROM employee_info where EMPLOYEE_ID = "+self.ID +";"
        cursor.execute(query)
        fetch = cursor.fetchall()
        print(fetch)
        w4.lineFirstname_3.setText(str(fetch[0][2]))
        w4.lineMid_3.setText(str(fetch[0][3]))
        w4.lineLastname_3.setText(str(fetch[0][1]))
        w4.linePhone_9.setText(str(fetch[0][4]))
        w4.linePhone_9.setText(str(fetch[0][4]))
        w4.dateEdit_3.setDate(QDate(fetch[0][5]))
        if fetch[0][6] == "MALE":
            w4.radioMale_3.setChecked(True)
        else:
            w4.radioFemale_3.setChecked(True)
        w4.lineDesignation_3.setText(str(fetch[0][7]))
        w4.lineSalaryg_3.setText(str(fetch[0][8]))
        w4.lineStepincre_3.setText(str(fetch[0][9]))
        w4.comboBox_3.setCurrentText(str(fetch[0][11]))
        w4.textEdit.setText(str(fetch[0][20]))
        w4.dateEdit.setDate(QDate(fetch[0][10]))
        w4.dateEdit_2.setDate(QDate(fetch[0][12]))
        w4.linePhone_8.setText(str(fetch[0][0]))
        firstID = str(fetch[0][0])
        w4.linePhone_7.setText(str(fetch[0][13]))
        w4.lineEmail_3.setText(str(fetch[0][14]))
        w4.lineTin_3.setText(str(fetch[0][15]))
        w4.lineGsis_3.setText(str(fetch[0][16]))
        w4.lineSss_3.setText(str(fetch[0][17]))
        w4.linePagibig_3.setText(str(fetch[0][18]))
        w4.linePhil_3.setText(str(fetch[0][19]))

        x = str(fetch[0][0])
        filename = x + '.png'
        pixmap = QPixmap(filename)
        w4.label.setPixmap(QPixmap(pixmap))
    #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    def showlistemploy(self):
        employtype = self.comboList.currentText()
        if employtype == 'TEACHING PERSONNEL':
            self.listE.clear()
            cur = conn.cursor()
            cur.execute("select LASTNAME,FIRSTNAME,MIDDLENAME from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        elif employtype == 'NON-TEACHING PERSONNEL':
            self.listE.clear()
            cur = conn.cursor()
            cur.execute("select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        elif employtype == 'PART TIME':
            self.listE.clear()
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        elif employtype == 'CONTRACT OF SERVICE':
            self.listE.clear()
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))

        elif employtype == 'SEPARATED':
            self.listE.clear()
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info where EMPLOYEE_STATUS = '" + employtype + "';")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        else:
            self.listE.clear()
            cur = conn.cursor()
            cur.execute(
                "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
            a = cur.fetchall()
            for i in range(len(a)):
                self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))

    def addemploybrowse(self):
        x = self.linePhone_8.text()
        if x == '':
            QMessageBox.warning(self,"Error","Please fill up the Employee's ID No.")
        else:
            lbl = QFileDialog.getOpenFileName(self,'Open File')
            print(lbl)
            if lbl[0] == '':
                print('waaaaaaaaaaaaa')
                self.label.clear()
            else:
                img = lbl[0]
                waa = Image.open(img)
                pixmap = QPixmap(img)
                print(pixmap)
                self.label.setPixmap(QPixmap(pixmap))
                
                filename = x+'.png'
                print(filename)
                waa.save(filename)
                self.browsedpic = filename
        '''options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        files, _ = QFileDialog.getOpenFileNames(self, "Browse Picture", "",
                                                "Images (*.png *.xpm *.jpg)", options=options)'''
        '''if files:
            self.browsedpic = files[0]
            self.label.setPixmap(QPixmap(self.browsedpic))'''
#UPPERCASEEEEEE~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    def auto_capitalfirstname(self, txt):
        upp_text = txt.upper()
        self.lineFirstname_3.setText(upp_text)

    def auto_capitalmiddlename(self, txt):
        upp_text = txt.upper()
        self.lineMid_3.setText(upp_text)
    def auto_capitallastname(self, txt):
        upp_text = txt.upper()
        self.lineLastname_3.setText(upp_text)

    def auto_capitaldesignation(self, txt):

        #regex = QtCore.QRegExp("^[a-zA-Z\s]*$")
        #validator = QtGui.QRegExpValidator(regex)
        upp_text = txt.upper()
        #self.lineDesignation_3.setValidator(validator)
        self.lineDesignation_3.setText(upp_text)

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~


#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~


    def infoexist(self,empID,lname,fname,mname):
        empID = empID
        lname = lname
        fname = fname
        mname = mname

        cur = conn.cursor()
        cur.execute("select EMPLOYEE_ID,LASTNAME,FIRSTNAME,MIDDLENAME from employee_info;")
        a = cur.fetchall()
        print(a)
        for i in range(len(a)):
            if str(a[i][0]) == str(empID) and str(a[i][1]) == str(lname) and str(a[i][2]) == str(fname) and str(
                    a[i][3]) == str(mname):
                return True

    def addemploy(self):
        fname = self.lineFirstname_3.text()
        mname = self.lineMid_3.text()
        lname = self.lineLastname_3.text()
        itemno = self.linePhone_9.text()
        bday = self.dateEdit_3.text()
        if self.radioMale_3.isChecked() == True:
            self.genders = "MALE"
        elif self.radioFemale_3.isChecked() == True:
            self.genders = "FEMALE"
        designation = self.lineDesignation_3.text()
        salarygrade = self.lineSalaryg_3.text()
        step = self.lineStepincre_3.text()
        empstatus = self.comboBox_3.currentText()

        dateoflastprom = self.dateEdit.text()
        dtt1 = dateoflastprom.split('/')
        day1 = int(dtt1[0]) + 3
        dtt1[0] = (day1)
        dayofpromotion = str(dtt1[0])+'/'+str(dtt1[1])+'/'+str(dtt1[2])
        print('dayofpromition '+dayofpromotion)
        firstdayofserv = self.dateEdit_2.text()
        dtt2 = firstdayofserv.split('/')
        dtt2[0] = int(dtt2[0])
        dtt2[1] = int(dtt2[1])
        dtt2[2] = int(dtt2[2])
        daytd = date.today()
        time1 = date(dtt2[0],dtt2[1],dtt2[2])
        print('waaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
        time2 = date.today()
        yearsofservice = time2.year - time1.year - ((time2.month,time2.day) < (time1.month,time1.day))
        print('yearsofsrvice '+str(yearsofservice))


        empID = self.linePhone_8.text()
        phonenum = self.linePhone_7.text()
        email = self.lineEmail_3.text()
        tin = self.lineTin_3.text()
        gsis = self.lineGsis_3.text()
        sss = self.lineSss_3.text()
        love = self.linePagibig_3.text()
        philhealth = self.linePhil_3.text()
        remarks = self.textEdit.toPlainText()
        lettersonlypattern = "^[a-zA-Z\s]*$"
        tendigitonlypattern = "^(\d{10})$"
        twelvedigitonlypattern = "^(\d{12})$"
        ninedigitonlypattern = "^(\d{9})$"
        elevendigitonlypattern = "^(\d{11})$"
        emailpattern = "[a-zA-Z0-9]+@[a-zA-Z]+\.[a-zA-Z]"
        phonenumpattern = "^(09)\d{9}$"
        emailcorrect = None
        twelveba = None
        twelvebaphilhealth = None
        nine = None
        ten = None
        eleven = None
        elevenphone = None


        if re.findall(phonenumpattern, phonenum):
            elevenphone =True
        else:
            elevenphone = False


        if re.findall(emailpattern, email):
            emailcorrect =True
        else:
            emailcorrect = False

        if re.findall(twelvedigitonlypattern, love):
            twelveba = True
        else:
            twelveba = False

        if re.findall(twelvedigitonlypattern, philhealth):
            twelvebaphilhealth = True
        else:
            twelvebaphilhealth = False

        if re.findall(ninedigitonlypattern, tin):
            nine = True
        else:
            nine = False

        if re.findall(tendigitonlypattern, sss):
            ten = True
        else:
            ten = False

        if re.findall(elevendigitonlypattern, gsis):
            eleven = True
        else:
            eleven = False

        # PHILIPPIANS 4:13 "I CAN DO ALL THINGS THROUGH CHRIST WHO STRENGTHENS ME"

        # JEREMIAH 29:11 FOR I KNOW THE PLANS I HAVE FOR YOU DECLARES THE LORD,PLANS FOR HOPE AND A FUTURE"
        # COLOSSIANS 3:23 WHATEVER YOU DO WORK HEARTILY AS FOR THE LORD NOT FOR MEN

        if fname == "" or lname == "" or itemno == "" or self.genders == None or  designation == "" or salarygrade == "" or step == "" or empstatus == "....." or empID == "" or phonenum == "" or email == "" or tin == "" or gsis == "" or sss =="" or love =="" or philhealth=="" or self.browsedpic == None:

            if fname =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's First Name")
            elif lname =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Last Name")

            elif itemno =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Item number")
            elif self.genders ==None:
                QMessageBox.warning(self,"Error","Please choose the Employee's Gender")
            elif designation =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Designation")
            elif salarygrade =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Salary Grade")
            elif step =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Step Increment")
            elif empstatus == ".....":
                QMessageBox.warning(self, "Error", "Please click and choose the Employee Status")
            elif empID =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's ID No.")
            elif phonenum =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Phone Number")
            elif email =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Email")
            elif tin =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's TIN number")
            elif gsis =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's GSIS")
            elif sss =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's SSS")
            elif philhealth =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's PhilHealth")
            elif self.browsedpic == None:
                QMessageBox.warning(self,"Error","Please choose the Employee's Picture")
        elif emailcorrect == False or  twelveba == False or twelvebaphilhealth == False or nine == False or ten==False or eleven == False or elevenphone == False:
            if emailcorrect == False:
                QMessageBox.warning(self,"Email Format Error","Please check the employee's email format\n It should be on this format:\nabcd122@gmail.com")
            elif twelveba == False:
                QMessageBox.warning(self, "PAGIBIG Input Error","PAGIBIG should be 12 digits")
            elif twelvebaphilhealth == False:
                QMessageBox.warning(self, "PHILHEALTH Input Error","PHILHEALTH should be 12 digits")
            elif nine == False:
                QMessageBox.warning(self, "TIN Input Error","TIN should be 9 digits")
            elif ten == False:
                QMessageBox.warning(self, "SSS Input Error","SSS should be 10 digits")
            elif eleven == False:
                QMessageBox.warning(self, "GSIS Input Error","GSIS should be 11 digits")

            elif elevenphone == False:
                QMessageBox.warning(self, "Phonenumber Input Error","Phone number should start with  '09' followed by other 9 digits")
        else:
            existing = w3.infoexist(empID,lname,fname,mname)
            print(existing)
            if existing:
                QMessageBox.warning(self,"Error","The employee named " + str(lname)+ ", "+ str(fname)+ " "+ str(mname)+"\nEmployee Number: "+ str(empID)+" \nis already existing on the Database")
            elif existing == None:
                try:
                    curAddEmp = conn.cursor()
                    queryAddEmp = "INSERT INTO employee_info(EMPLOYEE_ID, LASTNAME, FIRSTNAME, MIDDLENAME,ITEMNO, BIRTHDAY, GENDER, DESIGNATION, SALARY_GRADE, STEP_INCREMENT, DATE_OF_LAST_PROMOTION, FIRST_DAY_OF_SERVICE, PHONE_NUMBER, EMPLOYEE_STATUS, EMAIL, TIN, GSIS, SSS, PAGIBIG, PHILHEALTH, REMARKS, YEARS_OF_SERVICE, DAY_OF_PROMOTION) values( %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s);"
                    data = (empID, lname, fname, mname, itemno, bday, self.genders, designation, salarygrade, step, dateoflastprom,firstdayofserv, str(phonenum), empstatus, email, str(tin), str(gsis), str(sss), str(love), str(philhealth), str(remarks), str(yearsofservice), str(dayofpromotion))
                    curAddEmp.execute(queryAddEmp, data)
                    conn.commit()
                    QMessageBox.about(self, "Saved", "Information Saved to the Database Successfully")
                    self.listE.clear()
                    cur = conn.cursor()
                    cur.execute(
                        "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
                    a = cur.fetchall()
                    for i in range(len(a)):
                        self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
                    self.comboList.setCurrentText('ALL')
                    self.lineFirstname_3.setText("")
                    self.lineMid_3.setText("")
                    self.lineLastname_3.setText("")
                    self.linePhone_9.setText("")
                    self.dateEdit_3.setDate(date.today())
                    self.radioMale_3.setAutoExclusive(False)
                    self.radioMale_3.setChecked(False)
                    self.radioFemale_3.setAutoExclusive(False)
                    self.radioFemale_3.setChecked(False)
                    self.lineDesignation_3.setText("")
                    self.lineSalaryg_3.setText("")
                    self.lineStepincre_3.setText("")

                    self.comboBox_3.setCurrentText('.....')
                    self.dateEdit.setDate(date.today())
                    self.dateEdit_2.setDate(date.today())
                    self.linePhone_8.setText("")
                    self.linePhone_7.setText("")
                    self.lineEmail_3.setText("")
                    self.lineTin_3.setText("")
                    self.lineGsis_3.setText("")
                    self.lineSss_3.setText("")
                    self.linePagibig_3.setText("")
                    self.linePhil_3.setText("")
                    self.textEdit.setText("")

                    self.label.setText("            1X1 Picture*")
                    self.browsedpic == None

                except mysql.connector.Error as err:
                    if err.errno == 1062:
                        QMessageBox.warning(self,"Error","Employee ID No. Already Exists")
                    elif err.errno == 1048:
                        QMessageBox.warning(self, "Error", "Null")
                    elif err.errno == 1053:
                        QMessageBox.warning(self, "Error", "server shutdown")
                    else:
                        QMessageBox.warning(self, "Error", "ibapa")
    '''def addemploy(self):
        fname = self.lineFirstname_3.text()
        mname = self.lineMid_3.text()
        lname = self.lineLastname_3.text()
        itemno = self.linePhone_9.text()
        bday = self.dateEdit_3.text()
        if self.radioMale_3.isChecked() == True:
            self.genders = "MALE"
        elif self.radioFemale_3.isChecked() == True:
            self.genders = "FEMALE"
        designation = self.lineDesignation_3.text()
        salarygrade = self.lineSalaryg_3.text()
        step = self.lineStepincre_3.text()

        empstatus = self.comboBox_3.currentText()
        dateoflastprom = self.dateEdit.text()
        firstdayofserv = self.dateEdit_2.text()
        empID = self.linePhone_8.text()
        phonenum = self.linePhone_7.text()
        email = self.lineEmail_3.text()
        tin = self.lineTin_3.text()
        gsis = self.lineGsis_3.text()
        sss = self.lineSss_3.text()
        love = self.linePagibig_3.text()
        philhealth = self.linePhil_3.text()
        remarks = self.textEdit.toPlainText()
        lettersonlypattern = "^[a-zA-Z\s]*$"
        tendigitonlypattern = "^(\d{10})$"
        twelvedigitonlypattern = "^(\d{12})$"
        ninedigitonlypattern = "^(\d{9})$"
        elevendigitonlypattern = "^(\d{11})$"
        emailpattern = "[a-zA-Z0-9]+@[a-zA-Z]+\.[a-zA-Z]"
        phonenumpattern = "^(09)\d{9}$"
        emailcorrect = None
        twelveba = None
        twelvebaphilhealth = None
        nine = None
        ten = None
        eleven = None
        elevenphone = None


        if re.findall(phonenumpattern, phonenum):
            elevenphone =True
        else:
            elevenphone = False


        if re.findall(emailpattern, email):
            emailcorrect =True
        else:
            emailcorrect = False

        if re.findall(twelvedigitonlypattern, love):
            twelveba = True
        else:
            twelveba = False

        if re.findall(twelvedigitonlypattern, philhealth):
            twelvebaphilhealth = True
        else:
            twelvebaphilhealth = False

        if re.findall(ninedigitonlypattern, tin):
            nine = True
        else:
            nine = False

        if re.findall(tendigitonlypattern, sss):
            ten = True
        else:
            ten = False

        if re.findall(elevendigitonlypattern, gsis):
            eleven = True
        else:
            eleven = False

        # PHILIPPIANS 4:13 "I CAN DO ALL THINGS THROUGH CHRIST WHO STRENGTHENS ME"

        # JEREMIAH 29:11 FOR I KNOW THE PLANS I HAVE FOR YOU DECLARES THE LORD,PLANS FOR HOPE AND A FUTURE"
        # COLOSSIANS 3:23 WHATEVER YOU DO WORK HEARTILY AS FOR THE LORD NOT FOR MEN

        if fname == "" or lname == "" or itemno == "" or self.genders == None or  designation == "" or salarygrade == "" or step == "" or empstatus == "....." or empID == "" or phonenum == "" or email == "" or tin == "" or gsis == "" or sss =="" or love =="" or philhealth=="" or self.label == None:

            if fname =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's First Name")
            elif lname =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Last Name")

            elif itemno =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Item number")
            elif self.genders ==None:
                QMessageBox.warning(self,"Error","Please choose the Employee's Gender")
            elif designation =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Designation")
            elif salarygrade =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Salary Grade")
            elif step =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Step Increment")
            elif empstatus == ".....":
                QMessageBox.warning(self, "Error", "Please click and choose the Employee Status")
            elif empID =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's ID No.")
            elif phonenum =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Phone Number")
            elif email =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's Email")
            elif tin =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's TIN number")
            elif gsis =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's GSIS")
            elif sss =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's SSS")
            elif philhealth =="":
                QMessageBox.warning(self,"Error","Please fill up the Employee's PhilHealth")
            elif self.label == None:
                QMessageBox.warning(self,"Error","Please choose the Employee's Picture")
        elif emailcorrect == False or  twelveba == False or twelvebaphilhealth == False or nine == False or ten==False or eleven == False or elevenphone == False:
            if emailcorrect == False:
                QMessageBox.warning(self,"Email Format Error","Please check the employee's email format\n It should be on this format:\nabcd122@gmail.com")
            if twelveba == False:
                QMessageBox.warning(self, "PAGIBIG Input Error","PAGIBIG should be 12 digits")
            if twelvebaphilhealth == False:
                QMessageBox.warning(self, "PHILHEALTH Input Error","PHILHEALTH should be 12 digits")
            if nine == False:
                QMessageBox.warning(self, "TIN Input Error","TIN should be 9 digits")
            if ten == False:
                QMessageBox.warning(self, "SSS Input Error","SSS should be 10 digits")
            if eleven == False:
                QMessageBox.warning(self, "GSIS Input Error","GSIS should be 11 digits")

            if elevenphone == False:
                QMessageBox.warning(self, "Phonenumber Input Error","Phone number should start with  '09' followed by other 9 digits")
        else:
            existing = w3.infoexist(empID,lname,fname,mname)
            print(existing)
            if existing:
                QMessageBox.warning(self,"Error","The employee named " + str(lname)+ ", "+ str(fname)+ " "+ str(mname)+"\nEmployee Number: "+ str(empID)+" \nis already existing on the Database")
            elif existing == None:
                try:
                    curAddEmp = conn.cursor()
                    queryAddEmp = "INSERT INTO employee_info(EMPLOYEE_ID, LASTNAME, FIRSTNAME, MIDDLENAME,ITEMNO, BIRTHDAY, GENDER, DESIGNATION, SALARY_GRADE, STEP_INCREMENT, DATE_OF_LAST_PROMOTION, FIRST_DAY_OF_SERVICE, PHONE_NUMBER, EMPLOYEE_STATUS, EMAIL, TIN, GSIS, SSS, PAGIBIG, PHILHEALTH, REMARKS) values( %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s);"
                    data = (empID, lname, fname, mname, itemno, bday, self.genders, designation, salarygrade, step, dateoflastprom,firstdayofserv, str(phonenum), empstatus, email, str(tin), str(gsis), str(sss), str(love), str(philhealth), str(remarks))
                    curAddEmp.execute(queryAddEmp, data)
                    conn.commit()
                    QMessageBox.about(self, "Saved", "Information Saved to the Database Successfully")
                    self.listE.clear()
                    cur = conn.cursor()
                    cur.execute(
                        "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
                    a = cur.fetchall()
                    for i in range(len(a)):
                        self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
                    self.comboList.setCurrentText('ALL')

                    self.lineFirstname_3.setText("")
                    self.lineMid_3.setText("")
                    self.lineLastname_3.setText("")
                    self.linePhone_9.setText("")
                    self.dateEdit_3.setDate(date.today())
                    self.radioMale_3.setAutoExclusive(False)
                    self.radioMale_3.setChecked(False)
                    self.radioFemale_3.setAutoExclusive(False)
                    self.radioFemale_3.setChecked(False)
                    self.lineDesignation_3.setText("")
                    self.lineSalaryg_3.setText("")
                    self.lineStepincre_3.setText("")

                    self.comboBox_3.setCurrentText('.....')
                    self.dateEdit.setDate(date.today())
                    self.dateEdit_2.setDate(date.today())
                    self.linePhone_8.setText("")
                    self.linePhone_7.setText("")
                    self.lineEmail_3.setText("")
                    self.lineTin_3.setText("")
                    self.lineGsis_3.setText("")
                    self.lineSss_3.setText("")
                    self.linePagibig_3.setText("")
                    self.linePhil_3.setText("")
                    self.textEdit.setText("")
                    self.label.setText('2X2 Picture*')

                except mysql.connector.Error as err:
                    if err.errno == 1062:
                        QMessageBox.warning(self,"Error","Employee ID No. Already Exists")
                    elif err.errno == 1048:
                        QMessageBox.warning(self, "Error", "Null")
                    elif err.errno == 1053:
                        QMessageBox.warning(self, "Error", "server shutdown")
                    else:
                        QMessageBox.warning(self, "Error", "ibapa")'''
    def clicklistE(self):

        self.pic.show()
        self.textData.show()
        self.pushPrint.show()
        self.pushModify.show()
        self.textData.setText("")
        selectedrow = self.listE.currentItem().text()
        spltrow = selectedrow.split("  ")       # para maging list yung selected row sa list widget
        lastname = spltrow[0] # at dahil may kuwet sa last name ng nakuha nating string from list widget, from line 510 to  ay makikitang gumamit ako ng slice at pop method para matanggal yung kuwet sa string ng last name
        lastname = list(lastname)
        lastname.pop()
        lastname = "".join(lastname) # join method para yung last name ay maibalik natin into string
        spltrow[0] = lastname #binalik na natin sa list yung output ng pinaggagawa nating kagaguhan HAHHAAHA



        cursor = conn.cursor()
        query = "select EMPLOYEE_ID, LASTNAME, FIRSTNAME, MIDDLENAME,ITEMNO, BIRTHDAY, GENDER, DESIGNATION, SALARY_GRADE, STEP_INCREMENT, DATE_OF_LAST_PROMOTION, FIRST_DAY_OF_SERVICE, PHONE_NUMBER, EMPLOYEE_STATUS, EMAIL, TIN, GSIS, SSS, PAGIBIG, PHILHEALTH, REMARKS from employee_info where LASTNAME=%s and FIRSTNAME=%s and MIDDLENAME=%s;"
        data = (str(spltrow[0]), str(spltrow[1]), str(spltrow[2]))
        cursor.execute(query, data)
        fetch = cursor.fetchall()
        self.ID = str(fetch[0][0])

        imagelist = []
        HOME_DIR = os.path.dirname(__file__)
        x = str(fetch[0][0])
        print
        filename = x + '.png'
        print(filename)
        path = HOME_DIR
        images = [f for f in os.listdir(path)if os.path.splitext(f)[-1]=='.png']
        print(images)
        if images == []:
            QMessageBox.about(self, 'PICTURES', 'NO PICTURE')
        else:
            for image in images:
                im = Image.open(image)
                imim = im.filename
                imagelist.append(imim)

            for i in range(len(imagelist)):
                if imagelist[i] == filename:
                    print(imagelist)
                    pixmap = QPixmap(imagelist[i])
                    self.pic.setPixmap(QPixmap(pixmap))
                    imagelist.clear()
                    break
            else:
                QMessageBox.about(self,'PICTURES','NO SUCH PICTURE')

        self.textData.setText("SELECTED EMPLOYEE DATA SHEET \n \n \n EMPLOYEE ID: " + str(fetch[0][0]) + "\n \n"
                              + " LAST NAME: " + str(fetch[0][1]) + "\n \n FIRST NAME: " + str(
            fetch[0][2]) + "\n \n MIDDLE NAME: " + str(fetch[0][3])
                              + "\n \n ITEM NO.: " + str(fetch[0][4]) + "\n \n BIRTHDAY(yyyy-mm-dd): " + str(
            fetch[0][5]) + "\n \n GENDER: " + str(fetch[0][6])
                              + "\n \n DESIGNATION: " + str(fetch[0][7]) + "\n \n SALARY GRADE: " + str(
            fetch[0][8]) + "\n \n STEP INCREMENT: " + str(fetch[0][9])
                              + "\n \n DATE OF LAST PROMOTION(yyyy-mm-dd): " + str(
            fetch[0][10]) + "\n \n EMPLOYEE STATUS: " + str(
            fetch[0][11]) + "\n \n FIRST DAY OF SERVICE(yyyy-mm-dd): " + str(fetch[0][12])
                              + "\n \n PHONE NUMBER: " + str(fetch[0][13]) + "\n \n E-MAIL: " + str(
            fetch[0][14]) + "\n \n TIN: " + str(fetch[0][15]) + "\n \n GSIS: " + str(fetch[0][16])
                              + "\n \n SSS: " + str(fetch[0][17]) + "\n \n PAGIBIG: " + str(
            fetch[0][18]) + "\n \n PHILHEALTH: " + str(fetch[0][19]))

        if str(fetch[0][20]) == "":
            self.textData.append("\n REMARKS: None")
        else:
            self.textData.append("\n REMARKS: " + str(fetch[0][20]))

        


    def searchlist(self):
        self.listE.clear()
        search = self.lineSearch.text()
        cur = conn.cursor()
        cur.execute("select LASTNAME,FIRSTNAME,MIDDLENAME from employee_info where LASTNAME LIKE '%"+search+"%' or FIRSTNAME LIKE '%"+search+"%'  or MIDDLENAME LIKE '%"+search+"%';")
        a = cur.fetchall()
        for i in range(len(a)):
            self.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))

    def printinfoselectedemployee(self):

        selectedrow = self.listE.currentItem().text()
        spltrow = selectedrow.split("  ")
        lastname = spltrow[0]
        lastname = list(lastname)
        lastname.pop()                  #Kagaya lang ng ginawa ko sa method na clicklistE
        lastname = "".join(lastname)
        spltrow[0] = lastname
        cursor = conn.cursor()
        query = "select * from employee_info where LASTNAME=%s and FIRSTNAME=%s and MIDDLENAME=%s;"
        data = (str(spltrow[0]), str(spltrow[1]), str(spltrow[2]))
        cursor.execute(query, data)
        fetch = cursor.fetchall()
        print(fetch)
        print(fetch[0][2]+' waaaaaaaaaaaaaaaaaaaaa')
        fullname = fetch[0][2]+' '+fetch[0][3]+" "+ fetch[0][1]
        filename = str(fetch[0][0]) +'.png'
        doc = docx.Document()
        doc.add_heading(fullname+' - ID NO: '+str(fetch[0][0]),0).bold=True
        doc.add_picture(filename,width=Inches(1.1))
        doc.add_heading('INFORMATION',1).bold=True
        doc.add_paragraph("ITEM NO: " + str(fetch[0][4]))
        doc.add_paragraph("BIRTHDAY: " + str(fetch[0][5]))
        doc.add_paragraph("GENDER: " + str(fetch[0][6]))
        doc.add_paragraph("DESIGNATION: " + str(fetch[0][7]))
        doc.add_paragraph("EMPLOYEE STATUS: " + str(fetch[0][11]))
        doc.add_paragraph("FIRST DAY OF SERVICE: " + str(fetch[0][12]))
        doc.add_paragraph("PHONE NO: " + str(fetch[0][13]))
        doc.add_paragraph("EMAIL: " + str(fetch[0][14]))
        doc.add_paragraph('')
        doc.add_heading('BENEFITS DOCUMENTS',2).bold=True
        doc.add_paragraph("TIN: " + str(fetch[0][15]))
        doc.add_paragraph("GSIS: " + str(fetch[0][16]))
        doc.add_paragraph("SSS: " + str(fetch[0][17]))
        doc.add_paragraph("PAGIBIG: " + str(fetch[0][18]))
        doc.add_paragraph("PHILHEALTH: " + str(fetch[0][19]))
        doc.add_paragraph("REMARKS: " + str(fetch[0][20]))
        doc.save('Employee_info.docx')
        os.system('start Employee_info.docx')
        

    def clicklistnotif(self):
        self.frame_3.show()
        self.textviewnotif.show()
        self.pushupdatenotif.show()
        self.textviewnotif.clear()
        selectedrow = self.listnotif.currentItem().text()

        spltrow = selectedrow.split(",")
        print('waaaaaaaaaaaaaaaa')
        print(spltrow)
        lastname  =  spltrow[0]
        firstname = spltrow[1]
        middlename = spltrow[2]

        currlog=conn.cursor()
        querylog="SELECT DATE_OF_LAST_PROMOTION,DAY_OF_PROMOTION from employee_info where LASTNAME = %s and FIRSTNAME = %s and MIDDLENAME = %s;"
        datas = (lastname,firstname,middlename)
        currlog.execute(querylog,datas)
        a=currlog.fetchone()
        print(a)
        xxx = str(a[0])
        print(xxx)
        print('ditooooooooo')



        user = self.lineEdit_8.text()

        dateofnext = a[1]
        dateelastinword = xxx
        dateenextinword = str(dateofnext)

        self.textviewnotif.setHtml(
                                              "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                              "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                              "p, li { white-space: pre-wrap; }\n"
                                              "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:8.25pt; font-weight:400; font-style:normal;\">\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt;\">Dear </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+user+" </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#000000;\">:</span></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">    We would like to inform you that our system  detected an employee to be promote on </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+str(dateenextinword)+"</span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">.Please click the update button to update the date of last promotion of the employee. </span></p>\n"
                                              "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\"><br /></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">Employee Information:</span></p>\n"
                                              "<p align=\"justify\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\"><br /></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">Lastname: </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+lastname+"</span></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">Firstname: </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+firstname+"</span></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">Middlename: </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+middlename+"</span></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">Date of Last Promotion(yyyy-mm-dd): </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+str(dateelastinword)+"</span></p>\n"
                                              "<p align=\"justify\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600;\">Date of Next Promotion(yyyy-mm-dd): </span><span style=\" font-family:\'Arial Black\'; font-size:12pt; font-weight:600; color:#0055ff;\">"+str(dateenextinword)+"</span></p></body></html>")
    def updatenotif(self):
        selectedrow = self.listnotif.currentItem().text()
        spltrow = selectedrow.split(",")
        print('waaaaaaaaaaaaaaaa')
        print(spltrow)
        lastname  =  spltrow[0]
        firstname = spltrow[1]
        middlename = spltrow[2]

        currlog=conn.cursor()
        querylog="SELECT DATE_OF_LAST_PROMOTION,DAY_OF_PROMOTION from employee_info where LASTNAME = %s and FIRSTNAME = %s and MIDDLENAME = %s;"
        datas = (lastname,firstname,middlename)
        currlog.execute(querylog,datas)
        a=currlog.fetchone()
        print(a)
        xxx = str(a[0])

        dateofnext = str(a[1])
        dateoflastt = xxx
        print(xxx)
        print('ditooooooooo')

        print(lastname)
        print(firstname)
        print(middlename)
        print(dateofnext)
        print(dateoflastt)

        
        dateoflastprom = dateofnext
        dtt1 = dateoflastprom.split('/')
        dateofnextprom = str(dtt1[0])+'-'+str(dtt1[1])+'-'+str(dtt1[2])
        day1 = int(dtt1[0]) + 3
        dtt1[0] = (day1)
        dayofpromotion = str(dtt1[0])+'/'+str(dtt1[1])+'/'+str(dtt1[2])
        print('dayofpromotion '+dayofpromotion)
        print('dayofnextpromotion '+dateofnextprom)


        cursor = conn.cursor()
        query = "SELECT EMPLOYEE_ID FROM EMPLOYEE_INFO WHERE LASTNAME =%s AND FIRSTNAME = %s AND MIDDLENAME = %s;"
        data =(lastname,firstname,middlename)
        cursor.execute(query,data)
        fetchempid =cursor.fetchall()
        print(fetchempid[0][0])

        cursorupdate = conn.cursor()
        cursorupdate1 = conn.cursor()
        queryupdate = '''
            UPDATE EMPLOYEE_INFO
            SET DATE_OF_LAST_PROMOTION = %s
            WHERE EMPLOYEE_ID = %s
        '''
        queryupdate1 = '''
            UPDATE EMPLOYEE_INFO
            SET DAy_OF_PROMOTION = %s
            WHERE EMPLOYEE_ID = %s
        '''
        dataupdate = (dateofnext,str(fetchempid[0][0]))
        cursorupdate.execute(queryupdate,dataupdate)
        conn.commit()

        dataupdate1 = (dayofpromotion,str(fetchempid[0][0]))
        cursorupdate1.execute(queryupdate1,dataupdate1)
        conn.commit()

        QMessageBox.about(self,"Updated!","The Date of Last Promotion of employee named"+lastname+", "+ firstname + " "+ middlename +" has been changed")

        self.listnotif.clear()
        self.textviewnotif.clear()

        datee = date.today()
        datesss = str(datee)
        print('waaaaaaaaaaaaaaaa '+str(datee))
        x = datesss.split("-")

        currlog1=conn.cursor()
        querylog1="SELECT LASTNAME,FIRSTNAME,MIDDLENAME,DAY_OF_PROMOTION from employee_info;"
        currlog1.execute(querylog1)
        a1 =currlog1.fetchall()
        print(a1)

        pro_list = []
        for ii in range(len(a1)):
            date1 = a1[ii][3]
            date2 = str(date1)
            splt = date2.split('/')
            if x[0] == splt[0]:
                pro_list.append([a1[ii][0],a1[ii][1],a1[ii][2],a1[ii][3]])
        print('waaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
        print(pro_list)




        if len(pro_list):
            w3.listnotif.clear()
            for i in range(len(pro_list)):
                '''nextprom = self.fetchnotif[i][3]+ relativedelta(years=3)
                resultforprompt=abs(nextprom-date.today()).days
                print(resultforprompt)'''
                w3.listnotif.addItem(str(pro_list[i][0]) + "," + str(pro_list[i][1]) + "," + str(
                    pro_list[i][2]))
        else:
            w3.listnotif.clear()
            w3.listnotif.addItem("No notification so far")


        '''self.threedays = date.today() + timedelta(3)
        self.twodays = date.today() + timedelta(2)
        self.oneday = date.today() + timedelta(1)
        self.sub_3years3days = self.threedays + relativedelta(years=-3)
        self.sub_3years2days = self.twodays + relativedelta(years=-3)
        self.sub_3years1day = self.oneday + relativedelta(years=-3)
        print(self.sub_3years3days)
        print(self.sub_3years2days)
        print(self.sub_3years1day)
        self.cursornotif = conn.cursor()
        self.querynotif = ''SELECT LASTNAME,FIRSTNAME,MIDDLENAME,DATE_OF_LAST_PROMOTION 
        FROM EMPLOYEE_INFO WHERE DATE_OF_LAST_PROMOTION = %s or DATE_OF_LAST_PROMOTION = %s or DATE_OF_LAST_PROMOTION = %s;
        ''
        self.datanotif=(self.sub_3years3days,self.sub_3years2days,self.sub_3years1day)

        self.cursornotif.execute(self.querynotif,self.datanotif)
        self.fetchnotif = self.cursornotif.fetchall()
        if len(self.fetchnotif):
            w3.listnotif.clear()
            for i in range(len(self.fetchnotif)):

                nextprom = self.fetchnotif[i][3] + relativedelta(years=3)
                resultforprompt = abs(nextprom - date.today()).days
                print(resultforprompt)
                if resultforprompt == 3:
                    self.listnotif.addItem(str(self.fetchnotif[i][0]) + ", " + str(self.fetchnotif[i][1]) + " " + str(
                        self.fetchnotif[i][2]) + ".." + "just now")
                elif resultforprompt == 2:
                    self.listnotif.addItem(str(self.fetchnotif[i][0]) + ", " + str(self.fetchnotif[i][1]) + " " + str(
                        self.fetchnotif[i][2]) + ".." + "yesterday")
                elif resultforprompt == 1:
                    self.listnotif.addItem(str(self.fetchnotif[i][0]) + ", " + str(self.fetchnotif[i][1]) + " " + str(
                        self.fetchnotif[i][2]) + ".." + "2 days")'''


class newWindow4(Ui_MainWindowupd,QtWidgets.QMainWindow):
    def __init__(self):
        super(newWindow4, self).__init__()
        self.setupUi(self)
        self.setFixedSize(855, 566)
        self.label.setFixedSize(108, 108)
        self.pushButton.clicked.connect(self.backmain)
        self.pushbrowsepic.clicked.connect(self.modemploybrowse)
        self.lineFirstname_3.textChanged.connect(self.auto_capitalfirstname)
        self.lineMid_3.textChanged.connect(self.auto_capitalmiddlename)
        self.lineLastname_3.textChanged.connect(self.auto_capitallastname)
        self.lineDesignation_3.textChanged.connect(self.auto_capitaldesignation)
        self.pushAdd.clicked.connect(self.modifemploy)
        self.browsedpic = None
        regex = QtCore.QRegExp("^[a-zA-Z\s]*$")
        regnumnine = QtCore.QRegExp("^(\d{9})$")
        regnumten = QtCore.QRegExp("^(\d{10})$")
        regnumeleven = QtCore.QRegExp("^(\d{11})$")
        regnumtwelve = QtCore.QRegExp("^(\d{12})$")
        validatornine = QtGui.QRegExpValidator(regnumnine)
        validatorten = QtGui.QRegExpValidator(regnumten)
        validatoreleven = QtGui.QRegExpValidator(regnumeleven)
        validatortwelve = QtGui.QRegExpValidator(regnumtwelve)
        validator = QtGui.QRegExpValidator(regex)

        self.lineFirstname_3.setValidator(validator)
        self.lineMid_3.setValidator(validator)
        self.lineLastname_3.setValidator(validator)
        self.lineDesignation_3.setValidator(validator)
        self.lineStepincre_3.setValidator(QIntValidator())
        self.linePhone_9.setValidator(QIntValidator())
        self.lineSalaryg_3.setValidator(QIntValidator())
        self.lineStepincre_3.setValidator(QIntValidator())
        self.linePhone_8.setValidator(validatornine)
        self.linePhone_7.setValidator(validatoreleven)
        self.lineTin_3.setValidator(validatornine)
        self.lineGsis_3.setValidator(validatoreleven)
        self.lineSss_3.setValidator(validatorten)
        self.linePagibig_3.setValidator(validatortwelve)
        self.linePhil_3.setValidator(validatortwelve)

    def auto_capitalfirstname(self, txt):
        upp_text = txt.upper()
        self.lineFirstname_3.setText(upp_text)

    def auto_capitalmiddlename(self, txt):
        upp_text = txt.upper()
        self.lineMid_3.setText(upp_text)

    def auto_capitallastname(self, txt):
        upp_text = txt.upper()
        self.lineLastname_3.setText(upp_text)

    def auto_capitaldesignation(self, txt):
        # regex = QtCore.QRegExp("^[a-zA-Z\s]*$")
        # validator = QtGui.QRegExpValidator(regex)
        upp_text = txt.upper()
        # self.lineDesignation_3.setValidator(validator)
        self.lineDesignation_3.setText(upp_text)
    def modemploybrowse(self):
        x = self.linePhone_8.text()
        if x == '':
            QMessageBox.warning(self,"Error","Please fill up the Employee's ID No.")
        else:
            lbl = QFileDialog.getOpenFileName(self,'Open File')
            print(lbl)
            if lbl[0] == '':
                print('waaaaaaaaaaaaa')
                self.label.clear()
            else:
                img = lbl[0]
                waa = Image.open(img)
                pixmap = QPixmap(img)
                print(pixmap)
                self.label.setPixmap(QPixmap(pixmap))
                
                filename = x+'.png'
                print(filename)
                waa.save(filename)
            
        #if files:
        #   self.browsedpic = files[0]
        #   self.label.setPixmap(QPixmap(self.browsedpic))

    def backmain(self):
        w3.show()
        self.hide()
        w3.listE.clear()
        w3.pic.hide()
        w3.textData.hide()
        w3.pushModify.hide()
        w3.pushPrint.hide()
        w3.lineSearch.setText('')
        cur = conn.cursor()
        cur.execute(
            "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
        a = cur.fetchall()
        for i in range(len(a)):
            w3.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        w3.comboList.setCurrentText('ALL')

    def modifemploy(self):

        connn = mysql.connector.connect(host="localhost", user="root", passwd="admin", database="employeeinfodb")
        genders = None
        emailcorrect = None
        twelveba = None
        twelvebaphilhealth = None
        nine = None
        ten = None
        eleven = None
        elevenphone = None


        fname = self.lineFirstname_3.text()
        mname = self.lineMid_3.text()
        lname = self.lineLastname_3.text()
        itemno = self.linePhone_9.text()
        bday = self.dateEdit_3.text()
        if self.radioMale_3.isChecked() == True:
            genders = "MALE"
        elif self.radioFemale_3.isChecked() == True:
            genders = "FEMALE"
        designation = self.lineDesignation_3.text()
        salarygrade = self.lineSalaryg_3.text()
        step = self.lineStepincre_3.text()

        empstatus = self.comboBox_3.currentText()
        dateoflastprom = self.dateEdit.text()
        firstdayofserv = self.dateEdit_2.text()
        empID = self.linePhone_8.text()
        phonenum = self.linePhone_7.text()
        email = self.lineEmail_3.text()
        tin = self.lineTin_3.text()
        gsis = self.lineGsis_3.text()
        sss = self.lineSss_3.text()
        love = self.linePagibig_3.text()
        philhealth = self.linePhil_3.text()
        remarks = self.textEdit.toPlainText()


        lettersonlypattern = "^[a-zA-Z\s]*$"
        tendigitonlypattern = "^(\d{10})$"
        twelvedigitonlypattern = "^(\d{12})$"
        ninedigitonlypattern = "^(\d{9})$"
        elevendigitonlypattern = "^(\d{11})$"
        emailpattern = "[a-zA-Z0-9]+@[a-zA-Z]+\.[a-zA-Z]"
        phonenumpattern = "^(09)\d{9}$"

        if re.findall(phonenumpattern, phonenum):
            elevenphone = True
        else:
            elevenphone = False

        if re.findall(emailpattern, email):
            emailcorrect = True
        else:
            emailcorrect = False

        if re.findall(twelvedigitonlypattern, love):
            twelveba = True
        else:
            twelveba = False

        if re.findall(twelvedigitonlypattern, philhealth):
            twelvebaphilhealth = True
        else:
            twelvebaphilhealth = False

        if re.findall(ninedigitonlypattern, tin):
            nine = True
        else:
            nine = False

        if re.findall(tendigitonlypattern, sss):
            ten = True
        else:
            ten = False

        if re.findall(elevendigitonlypattern, gsis):
            eleven = True
        else:
            eleven = False

        # PHILIPPIANS 4:13 "I CAN DO ALL THINGS THROUGH CHRIST WHO STRENGTHENS ME"

        # JEREMIAH 29:11 FOR I KNOW THE PLANS I HAVE FOR YOU DECLARES THE LORD,PLANS FOR HOPE AND A FUTURE"
        # COLOSSIANS 3:23 WHATEVER YOU DO WORK HEARTILY AS FOR THE LORD NOT FOR MEN

        if fname == "" or lname == "" or itemno == "" or genders == None or designation == "" or salarygrade == "" or step == "" or empstatus == "....." or empID == "" or phonenum == "" or email == "" or tin == "" or gsis == "" or sss == "" or love == "" or philhealth == "":

            if fname == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's First Name")
            if lname == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Last Name")

            if itemno == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Item number")
            if genders == None:
                QMessageBox.warning(self, "Error", "Please choose the Employee's Gender")
            if designation == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Designation")
            if salarygrade == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Salary Grade")
            if step == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Step Increment")
            if empstatus == ".....":
                QMessageBox.warning(self, "Error", "Please click and choose the Employee Status")
            if empID == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's ID No.")
            if phonenum == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Phone Number")
            if email == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's Email")
            if tin == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's TIN number")
            if gsis == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's GSIS")
            if sss == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's SSS")
            if philhealth == "":
                QMessageBox.warning(self, "Error", "Please fill up the Employee's PhilHealth")
        elif emailcorrect == False or twelveba == False or twelvebaphilhealth == False or nine == False or ten == False or eleven == False or elevenphone == False:
            if emailcorrect == False:
                QMessageBox.warning(self, "Email Format Error",
                                    "Please check the employee's email format\n It should be on this format:\nabcd122@gmail.com")
            if twelveba == False:
                QMessageBox.warning(self, "PAGIBIG Input Error", "PAGIBIG should be 12 digits")
            if twelvebaphilhealth == False:
                QMessageBox.warning(self, "PHILHEALTH Input Error", "PHILHEALTH should be 12 digits")
            if nine == False:
                QMessageBox.warning(self, "TIN Input Error", "TIN should be 9 digits")
            if ten == False:
                QMessageBox.warning(self, "SSS Input Error", "SSS should be 10 digits")
            if eleven == False:
                QMessageBox.warning(self, "GSIS Input Error", "GSIS should be 11 digits")

            if elevenphone == False:
                QMessageBox.warning(self, "Phonenumber Input Error",
                                    "Phone number should start with  '09' followed by other 9 digits")
        else:
            try:
                cursor1 = conn.cursor()
                qwe = 'DELETE FROM employee_info WHERE employee_id = ' + empID + ';'
                cursor1.execute(qwe)
                conn.commit()
                curAddEmp = connn.cursor()
                queryAddEmp = "INSERT INTO employee_info(EMPLOYEE_ID, LASTNAME, FIRSTNAME, MIDDLENAME,ITEMNO, BIRTHDAY, GENDER, DESIGNATION, SALARY_GRADE, STEP_INCREMENT, DATE_OF_LAST_PROMOTION, FIRST_DAY_OF_SERVICE, PHONE_NUMBER, EMPLOYEE_STATUS, EMAIL, TIN, GSIS, SSS, PAGIBIG, PHILHEALTH, REMARKS) values( %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,%s);"
                data = (
                empID, lname, fname, mname, itemno, bday, genders, designation, salarygrade, step, dateoflastprom,
                firstdayofserv, str(phonenum), empstatus, email, str(tin), str(gsis), str(sss), str(love),
                str(philhealth), str(remarks))
                curAddEmp.execute(queryAddEmp, data)
                connn.commit()
                QMessageBox.about(self, "Modified", "Information Changed to the Database Successfully")

                self.lineFirstname_3.setText("")
                self.lineMid_3.setText("")
                self.lineLastname_3.setText("")
                self.linePhone_9.setText("")
                self.dateEdit_3.setDate(date.today())
                self.radioMale_3.setAutoExclusive(False)
                self.radioMale_3.setChecked(False)
                self.radioFemale_3.setAutoExclusive(False)
                self.radioFemale_3.setChecked(False)
                self.lineDesignation_3.setText("")
                self.lineSalaryg_3.setText("")
                self.lineStepincre_3.setText("")

                self.comboBox_3.setCurrentText('.....')
                self.dateEdit.setDate(date.today())
                self.dateEdit_2.setDate(date.today())
                self.linePhone_8.setText("")
                self.linePhone_7.setText("")
                self.lineEmail_3.setText("")
                self.lineTin_3.setText("")
                self.lineGsis_3.setText("")
                self.lineSss_3.setText("")
                self.linePagibig_3.setText("")
                self.linePhil_3.setText("")
                self.textEdit.setText("")
                self.label.setText('2X2 Picture*')


            except mysql.connector.Error as err:
                if err.errno == 1062:
                    QMessageBox.warning(self, "Error", "Employee ID No. Already Exists")
                elif err.errno == 1048:
                    QMessageBox.warning(self, "Error", "Null")
                elif err.errno == 1053:
                    QMessageBox.warning(self, "Error", "server shutdown")
                else:
                    QMessageBox.warning(self, "Error", "ibapa")
    def closeEvent(self, event):
        w3.show()
        event.accept()
        w3.listE.clear()
        cur = conn.cursor()
        cur.execute(
            "select LASTNAME,FIRSTNAME,MIDDLENAME,EMPLOYEE_STATUS from employee_info;")
        a = cur.fetchall()
        for i in range(len(a)):
            w3.listE.addItem(str(a[i][0]) + "," + "  " + str(a[i][1]) + "  " + str(a[i][2]))
        w3.comboList.setCurrentText('ALL')


        
if __name__ == "__main__":
    apps = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    w1 = newWindow1()
    w1.show()
    w2 = newWindow2()
    w2.hide()
    w3 = newWindow3()
    w3.hide()
    w4 = newWindow4()
    w4.hide()
    apps.exec_()